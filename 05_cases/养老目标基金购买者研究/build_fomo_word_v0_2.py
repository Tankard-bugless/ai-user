"""Generate the Word copy synchronized with the FOMO Tencent Survey deployment."""

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "output" / "doc"
SOURCE = OUTPUT_DIR / "养老目标基金购买与持有情况调查问卷_v0.2.docx"
TARGET = OUTPUT_DIR / "养老目标基金购买与持有情况调查问卷_FOMO_v0.2.docx"
SURVEY_URL = "https://wj.qq.com/s2/27394126/fb35"


def insert_after(anchor: Paragraph, text: str, style: str | None = None) -> Paragraph:
    node = OxmlElement("w:p")
    anchor._p.addnext(node)
    paragraph = Paragraph(node, anchor._parent)
    if style:
        paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    document = Document(SOURCE)
    document.core_properties.title = "养老目标基金购买与持有情况调查问卷 FOMO v0.2"
    document.core_properties.subject = "与 FOMO 腾讯问卷部署同步的 Word 版本"
    document.core_properties.comments = (
        "参与者正文与 FOMO 腾讯问卷 27394126 同版；"
        "平台自定义跳转因团队权限尚未启用。"
    )

    appendix = next(
        (
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text.strip() == "内部配置附录"
        ),
        None,
    )
    if appendix is None:
        raise ValueError("内部配置附录 heading not found")

    note = (
        "FOMO 团队同步记录（2026-07-24）：线上问卷 "
        f"{SURVEY_URL} 的参与者正文已核对为 10 页、35 道可作答题和 213 个选项。"
        "C0 不同意结束、S1 未购买结束、Q22 跳至 Q24 三条逻辑因当前团队权限尚未在线启用；"
        "本 Word 保留预期跳转作为设计与数据口径契约，解决权限或批准兼容方案前不得标记为正式投放版。"
    )
    insert_after(appendix, note, "Instruction")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(TARGET)

    check = Document(TARGET)
    participant_text: list[str] = []
    for paragraph in check.paragraphs:
        if paragraph.text.strip() == "内部配置附录":
            break
        participant_text.append(paragraph.text.strip())

    expected_ids = ["C0.", "S1."] + [f"Q{i}." for i in range(1, 10)]
    expected_ids += ["Q9A."] + [f"Q{i}." for i in range(10, 33)]
    question_count = sum(
        1
        for text in participant_text
        if any(text.startswith(prefix) for prefix in expected_ids)
    )
    option_count = sum(1 for text in participant_text if text.startswith("☐"))
    if question_count != 35:
        raise ValueError(f"Unexpected question count: {question_count}")
    if option_count != 213:
        raise ValueError(f"Unexpected option count: {option_count}")
    if SURVEY_URL not in "\n".join(paragraph.text for paragraph in check.paragraphs):
        raise ValueError("FOMO deployment note missing")

    print(
        f"created={TARGET}; questions={question_count}; "
        f"options={option_count}; tables={len(check.tables)}"
    )


if __name__ == "__main__":
    main()
