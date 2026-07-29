"""Build the OTF questionnaire v0.2 and its platform-neutral InstrumentSpec.

The approved v0.1 Word file is used as a visual template. The script writes a
new v0.2 file and never overwrites v0.1. Participant-facing questions and the
machine-readable InstrumentSpec are derived from the same document version.
"""

from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
CASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = ROOT / "output" / "doc"
SOURCE = OUTPUT_DIR / "养老目标基金购买与持有情况调查问卷_v0.1.docx"
TARGET = OUTPUT_DIR / "养老目标基金购买与持有情况调查问卷_v0.2.docx"
INSTRUMENT_TARGET = CASE_DIR / "instrument-spec.v0.2.0.json"
IMA_URL = (
    "https://ima.qq.com/wiki/?shareId="
    "3f82c28e1844208f0d4489f0ef3057c012f070cdce772a8ee050212c899234fc"
)

EXPECTED_IDS = (
    ["C0", "S1"]
    + [f"Q{i}" for i in range(1, 10)]
    + ["Q9A"]
    + [f"Q{i}" for i in range(10, 33)]
)
OPTIONAL_IDS = {"Q11", "Q29"}
MULTI_IDS = {"Q6", "Q7", "Q8", "Q21", "Q23", "Q25", "Q26", "Q27", "Q28"}
TEXT_SHORT_IDS = {"Q11"}
TEXT_LONG_IDS = {"Q29"}

RQ_MAP: dict[str, list[str]] = {
    "Q1": ["RQ-005"],
    "Q2": ["RQ-004"],
    "Q3": ["RQ-001"],
    "Q4": ["RQ-001"],
    "Q5": ["RQ-001"],
    "Q6": ["RQ-001", "RQ-004"],
    "Q7": ["RQ-001", "RQ-004"],
    "Q8": ["RQ-001", "RQ-004"],
    "Q9": ["RQ-001", "RQ-004"],
    "Q9A": ["RQ-004"],
    "Q10": ["RQ-002"],
    "Q11": ["RQ-002"],
    "Q12": ["RQ-002"],
    "Q13": ["RQ-003"],
    "Q14": ["RQ-003"],
    "Q15": ["RQ-003"],
    "Q16": ["RQ-003"],
    "Q17": ["RQ-002"],
    "Q18": ["RQ-002"],
    "Q19": ["RQ-003"],
    "Q20": ["RQ-003"],
    "Q21": ["RQ-005"],
    "Q22": ["RQ-005"],
    "Q23": ["RQ-005"],
    "Q24": ["RQ-005"],
    "Q25": ["RQ-005"],
    "Q26": ["RQ-006"],
    "Q27": ["RQ-006"],
    "Q28": ["RQ-006"],
    "Q29": ["RQ-006"],
    "Q30": ["RQ-004", "RQ-005"],
    "Q31": ["RQ-004"],
    "Q32": ["RQ-001", "RQ-002", "RQ-003", "RQ-006"],
}

HYP_MAP: dict[str, list[str]] = {
    **{qid: ["HYP-003"] for qid in ["Q3", "Q4", "Q5", "Q6", "Q7", "Q8", "Q9", "Q9A"]},
    **{qid: ["HYP-001"] for qid in ["Q10", "Q11", "Q12", "Q17", "Q18"]},
    **{qid: ["HYP-002"] for qid in ["Q13", "Q14", "Q15", "Q16", "Q19", "Q20"]},
    **{qid: ["HYP-004"] for qid in ["Q1", "Q21", "Q22", "Q23", "Q24", "Q25", "Q26", "Q27", "Q28", "Q29", "Q30"]},
}

FACT_MAP = {
    "S1": ["FF-FOF-003", "FF-FOF-009"],
    "Q2": ["FF-FOF-010"],
    "Q10": ["FF-FOF-003", "FF-FOF-004", "FF-FOF-005"],
    "Q17": ["FF-FOF-004", "FF-FOF-007"],
    "Q18": ["FF-FOF-005", "FF-FOF-007"],
    "Q19": ["FF-FOF-006"],
    "Q20": ["FF-FOF-007"],
}

SECTION_PURPOSE = {
    "填写说明": "取得研究同意并确认购买经历。",
    "一、最近一次购买经历": "记录最近一次购买背景和当前持有状态。",
    "二、当时是怎么选择的": "区分信息接触、认真比较、主要影响因素和主观清晰度。",
    "三、您对产品类型的了解": "测量产品类型自报与主观了解。",
    "四、您是否看过下滑曲线": "测量下滑曲线的知晓、查看和使用行为。",
    "五、概念理解": "使用经事实审核的知识题测量客观理解。",
    "六、购买后的持有情况": "记录持有行为、波动应对和信息不确定。",
    "七、您希望获得哪些信息": "收集信息主题、形式和时点偏好。",
    "八、简单背景信息": "收集预设研究变量和唯一基础分群年龄段。",
}


@dataclass
class Question:
    qid: str
    prompt: str
    section: str
    instruction: str = ""
    skip_note: str = ""
    options: list[str] = field(default_factory=list)


def find_paragraph(document: Document, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def insert_after(anchor: Paragraph, text: str, style: str) -> Paragraph:
    node = OxmlElement("w:p")
    anchor._p.addnext(node)
    paragraph = Paragraph(node, anchor._parent)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def replace_version_text(document: Document) -> None:
    for paragraph in document.paragraphs:
        if "V0.1" in paragraph.text:
            for run in paragraph.runs:
                if "V0.1" in run.text:
                    run.text = run.text.replace("V0.1", "V0.2")
    for section in document.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                for run in paragraph.runs:
                    if "V0.1" in run.text:
                        run.text = run.text.replace("V0.1", "V0.2")


def add_question_block(anchor: Paragraph, parts: list[tuple[str, str]]) -> Paragraph:
    current = anchor
    for style, text in parts:
        current = insert_after(current, text, style)
    return current


def find_row(table, first_cell_text: str):
    for row in table.rows:
        if row.cells[0].text.strip() == first_cell_text:
            return row
    raise ValueError(f"Table row not found: {first_cell_text}")


def set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def insert_row_after(table, first_cell_text: str, values: list[str]) -> None:
    target_row = find_row(table, first_cell_text)
    new_row = table.add_row()
    for cell, value in zip(new_row.cells, values, strict=True):
        set_cell(cell, value)
    target_row._tr.addnext(new_row._tr)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(document: Document, headers: list[str], rows: list[list[str]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Normal Table"
    table.autofit = True
    for cell, text in zip(table.rows[0].cells, headers, strict=True):
        set_cell(cell, text, bold=True)
    repeat_table_header(table.rows[0])
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values, strict=True):
            set_cell(cell, value)
    return table


def update_word() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    document = Document(SOURCE)
    replace_version_text(document)
    document.core_properties.title = "养老目标基金购买与持有情况调查问卷 v0.2"
    document.core_properties.subject = "问卷正文与数据有效性契约同版"
    document.core_properties.comments = (
        "v0.2 新增购买经历甄别、五点信息清晰度、年龄段和统一答卷/数据集状态。"
    )

    consent_no = find_paragraph(document, "☐  不同意")
    add_question_block(
        consent_no,
        [
            ("Question", "S1. 在填写本问卷之前，您是否实际购买过养老目标基金？"),
            (
                "Instruction",
                "单选；仅开立个人养老金账户、缴存资金或购买储蓄、保险、理财，不算购买过养老目标基金",
            ),
            ("Option", "☐  是，目前持有或过去曾经持有过"),
            ("Option", "☐  没有"),
            ("Option", "☐  不确定自己买的是不是养老目标基金"),
            (
                "SkipNote",
                "跳转提示：选择“没有”时结束问卷；选择“不确定”可继续作答，数据返回后只进入人工复核，不自动删除。",
            ),
        ],
    )

    section_three = find_paragraph(document, "三、您对产品类型的了解")
    paragraph_list = document.paragraphs
    section_three_index = next(
        index for index, paragraph in enumerate(paragraph_list) if paragraph._p is section_three._p
    )
    q9_anchor = paragraph_list[section_three_index - 1]
    add_question_block(
        q9_anchor,
        [
            ("Question", "Q9A. 总体来说，您当时觉得购买前看到的产品信息有多清楚？"),
            ("Instruction", "五点评分；1 表示非常不清楚，5 表示非常清楚；“记不清”单独列出"),
            ("Option", "☐  1—非常不清楚"),
            ("Option", "☐  2—比较不清楚"),
            ("Option", "☐  3—一般"),
            ("Option", "☐  4—比较清楚"),
            ("Option", "☐  5—非常清楚"),
            ("Option", "☐  记不清"),
        ],
    )

    completion_heading = find_paragraph(document, "问卷结束")
    paragraph_list = document.paragraphs
    completion_index = next(
        index
        for index, paragraph in enumerate(paragraph_list)
        if paragraph._p is completion_heading._p
    )
    q31_anchor = paragraph_list[completion_index - 1]
    add_question_block(
        q31_anchor,
        [
            ("Question", "Q32. 您的年龄段是？"),
            ("Instruction", "单选；仅用于本次匿名样本的方向性分组比较"),
            ("Option", "☐  18—29 岁"),
            ("Option", "☐  30—39 岁"),
            ("Option", "☐  40—49 岁"),
            ("Option", "☐  50—59 岁"),
            ("Option", "☐  60 岁及以上"),
            ("Option", "☐  不愿回答"),
        ],
    )

    # Existing appendix tables remain authoritative for questionnaire operations.
    jump_table = document.tables[0]
    insert_row_after(
        jump_table,
        "C0",
        ["S1", "选择“没有”时结束；选择“不确定”继续并标记 ELIGIBILITY_UNCERTAIN。"],
    )

    variable_table = document.tables[1]
    insert_row_after(
        variable_table,
        "C0",
        ["S1", "purchase_eligibility", "单选", "-", "必答；没有→结束，不确定→继续并复核。"],
    )
    insert_row_after(
        variable_table,
        "Q9",
        ["Q9A", "purchase_information_clarity_5", "五点评分", "RQ-004", "必答；记不清为独立有效值。"],
    )
    insert_row_after(
        variable_table,
        "Q31",
        [
            "Q32",
            "age_group",
            "单选",
            "RQ-001/002/003/006",
            "必答；不愿回答为有效值，但不进入具体年龄组比较。",
        ],
    )

    report_table = document.tables[3]
    report_updates = {
        "RQ-001": ("Q3-Q9、Q32", "购买路径与不同年龄段的信息渠道、决策角色"),
        "RQ-002": ("Q10-Q12、Q17-Q18、Q32", "产品类型识别、客观理解及年龄方向差异"),
        "RQ-003": ("Q13-Q16、Q19-Q20、Q32", "下滑曲线的知晓—查看—使用—理解及年龄方向差异"),
        "RQ-004": ("Q2、Q7-Q9、Q9A", "决策因素与购买前信息清晰度"),
        "RQ-005": ("Q1、Q21-Q25、Q30", "持有行为、波动应对和不确定"),
        "RQ-006": ("Q26-Q29、Q32", "信息主题、形式、时点及年龄方向差异"),
    }
    for rq, (evidence, output) in report_updates.items():
        row = find_row(report_table, rq)
        set_cell(row.cells[1], evidence)
        set_cell(row.cells[2], output)

    # Update existing checks without renumbering the frozen v0.1 document.
    check_1 = find_paragraph(
        document, "检查 1  全部题号、变量名和选项值唯一；导出字段与变量表一致。"
    )
    check_1.text = (
        "检查 1  全部题号、变量名和选项值唯一；导出字段与 InstrumentSpec v0.2.0 一致。"
    )
    check_2 = find_paragraph(
        document, "检查 2  C0、Q22→Q23 跳转可达，互斥选项和最多选择数量已配置。"
    )
    check_2.text = (
        "检查 2  C0、S1、Q22→Q23 跳转可达；互斥选项和最多选择数量已配置。"
    )
    check_4 = find_paragraph(
        document, "检查 4  Q17-Q20 的答案与经审核的产品事实一致。"
    )
    check_4.text = (
        "检查 4  S1、Q2、Q10、Q17-Q20 的金融陈述与 FIN-FACT-FOF@0.1.0 一致。"
    )

    heading = document.add_heading("F. 数据结构与有效性状态", level=2)
    heading.paragraph_format.page_break_before = True
    document.add_paragraph(
        "本附录与问卷正文同版。平台导出后不得重新发明口径；先核对版本和字段，"
        "再按以下值语义和状态规则处理。知识题答错、选择不确定、年龄不愿回答和开放题留空，"
        "都不能单独判为无效答卷。"
    )
    add_table(
        document,
        ["值状态", "含义", "分析处理"],
        [
            ["SUBSTANTIVE", "正常选项、文本或数值", "进入适用题目分母"],
            ["UNKNOWN", "不知道、不确定、记不清", "作为有效结果保留；知识题可计未答对"],
            ["NOT_APPLICABLE", "题目不适用", "不进入该题适用分母"],
            ["PREFER_NOT_TO_SAY", "明确不愿回答", "保留为独立类别，不据此排除答卷"],
            ["STRUCTURAL_MISSING", "因合法跳转未展示", "不是漏答，不进入该题分母"],
            ["ITEM_MISSING", "应展示但没有答案", "计入题目缺失，不自动排除整份答卷"],
        ],
    )
    document.add_paragraph()
    add_table(
        document,
        ["答卷状态", "含义", "主分析处理"],
        [
            ["UNASSESSED", "已导入、尚未执行规则", "不得进入冻结数据集"],
            ["VALID", "满足同意和目标人群要求，无未解决问题", "纳入"],
            ["REVIEW_REQUIRED", "有可疑信号但不足以自动判无效", "人工复核后转 VALID 或 EXCLUDED"],
            ["EXCLUDED", "命中硬排除或经人工确认排除", "不纳入，保留原因和决定记录"],
        ],
    )
    document.add_paragraph()
    add_table(
        document,
        ["类型", "原因码", "当前 Demo 规则"],
        [
            ["硬排除", "CONSENT_NOT_GIVEN", "C0 不同意或缺失"],
            ["硬排除", "INELIGIBLE", "S1 明确选择没有购买过"],
            ["硬排除", "NO_SUBSTANTIVE_DATA", "除同意/甄别外核心模块均无有效信息"],
            ["硬排除", "SCHEMA_OR_VERSION_ERROR", "无法对应批准问卷版本且不能可靠修复"],
            ["硬排除", "CONFIRMED_DUPLICATE", "人工确认重复，按规则保留一份"],
            ["需复核", "ELIGIBILITY_UNCERTAIN", "S1 不确定自己是否买过养老目标基金"],
            ["需复核", "SPEEDING_SUSPECTED", "至少 30 份后，时长低于当批中位数三分之一"],
            ["需复核", "LOGIC_CONFLICT", "关键事实或路由字段出现无法解释的冲突"],
            ["需复核", "DUPLICATE_SUSPECTED", "匿名记录疑似重复但尚未确认"],
            ["需复核", "HIGH_ITEM_MISSINGNESS", "应展示核心题的 ITEM_MISSING 超过 20%"],
        ],
    )
    document.add_paragraph()
    add_table(
        document,
        ["数据集状态/输出", "规则"],
        [
            ["ANALYSIS_READY", "所有答卷状态已解决，字段和分母完整，无重要限制"],
            ["ANALYSIS_READY_WITH_LIMITS", "可分析但存在样本来源集中、子组过小或缺失等非阻断限制"],
            ["BLOCKED", "仍有 UNASSESSED/REVIEW_REQUIRED，或版本、字段、核心逻辑无法核对"],
            ["年龄 n<10", "不单独展示比例；必要时合并相邻年龄段或标记样本不足"],
            ["年龄 10≤n<20", "可展示人数和比例，但标注“小样本，仅供方向观察”"],
            ["年龄 n≥20", "可作描述性比例比较，仍不作总体、显著性或因果推断"],
        ],
    )

    document.add_heading("G. v0.2 发布前检查", level=2)
    for index, text in enumerate(
        [
            "问卷正文、InstrumentSpec、腾讯问卷 DSL 和变量表的版本及题号完全一致。",
            "C0 不同意与 S1 没有购买均结束；S1 不确定继续，并输出 ELIGIBILITY_UNCERTAIN。",
            "Q9A 的 1—5 方向一致，“记不清”不占用量表中点。",
            "Q32 只收年龄区间，不收生日；“不愿回答”不导致答卷无效。",
            "S1、Q2、Q10、Q17—Q20 已通过 FIN-FACT-FOF@0.1.0 事实复核。",
            "所有 REVIEW_REQUIRED 在数据冻结前已由研究人员解决。",
            "报告中的样本数、题目分母、缺失、排除和年龄子组状态可回溯。",
        ],
        start=1,
    ):
        document.add_paragraph(f"检查 {index}  {text}")

    TARGET.parent.mkdir(parents=True, exist_ok=True)
    document.save(TARGET)


def parse_questionnaire(path: Path) -> tuple[list[Question], list[str]]:
    document = Document(path)
    questions: list[Question] = []
    section_order: list[str] = []
    current_section = ""
    current: Question | None = None
    participant_started = False

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        if text == "内部配置附录" or text == "问卷结束":
            break
        if style == "Heading 1" and text == "填写说明":
            participant_started = True
        if not participant_started:
            continue
        if style == "Heading 1":
            current_section = text
            section_order.append(text)
            current = None
            continue
        if style == "Question":
            match = re.match(r"^(C0|S1|Q\d+(?:A)?)\.\s*(.+)$", text)
            if not match:
                raise ValueError(f"Cannot parse question: {text}")
            current = Question(match.group(1), match.group(2), current_section)
            questions.append(current)
            continue
        if current is None:
            continue
        if style == "Instruction":
            current.instruction = text
        elif style == "Option":
            current.options.append(re.sub(r"^☐\s*", "", text))
        elif style == "SkipNote":
            current.skip_note = re.sub(r"^跳转提示：\s*", "", text)

    actual_ids = [question.qid for question in questions]
    if actual_ids != EXPECTED_IDS:
        raise ValueError(f"Unexpected question order: {actual_ids}")
    option_count = sum(len(question.options) for question in questions)
    if option_count != 213:
        raise ValueError(f"Unexpected option count: {option_count}")
    return questions, section_order


def variable_map(path: Path) -> dict[str, str]:
    document = Document(path)
    table = document.tables[1]
    mapping = {}
    for row in table.rows[1:]:
        qid = row.cells[0].text.strip()
        variable = row.cells[1].text.strip()
        if qid and variable:
            mapping[qid] = variable
    missing = set(EXPECTED_IDS) - set(mapping)
    if missing:
        raise ValueError(f"Variables missing for: {sorted(missing)}")
    return mapping


def normalized_value(qid: str, label: str, index: int):
    if qid == "Q9A":
        match = re.match(r"([1-5])", label)
        return int(match.group(1)) if match else 98
    if "不愿回答" in label and "不确定" not in label:
        return "prefer_not"
    if "没有经历过或记不清" in label:
        return "not_applicable"
    if label in {"不确定", "记不清"} or "不确定自己" in label:
        return "unknown"
    if "不确定或不愿回答" in label:
        return "unknown_or_prefer_not"
    return f"o{index:02d}"


def value_semantic(label: str) -> str | None:
    if "不愿回答" in label and "不确定" not in label:
        return "PREFER_NOT_TO_SAY"
    if "没有经历过或记不清" in label:
        return "NOT_APPLICABLE"
    if (
        label in {"不确定", "记不清"}
        or "不确定自己" in label
        or "不确定或不愿回答" in label
    ):
        return "UNKNOWN"
    return None


def item_type(question: Question) -> str:
    if question.qid == "C0":
        return "CONSENT"
    if question.qid == "Q9A":
        return "LIKERT"
    if question.qid in MULTI_IDS:
        return "MULTI_SELECT"
    if question.qid in TEXT_SHORT_IDS:
        return "TEXT_SHORT"
    if question.qid in TEXT_LONG_IDS:
        return "TEXT_LONG"
    return "SINGLE_SELECT"


def exclusive_option(qid: str, label: str) -> bool:
    rules = {
        "Q6": {"记不清"},
        "Q7": {"没有仔细看", "记不清"},
        "Q8": {"主要听取推荐，没有认真比较", "记不清"},
        "Q21": {"一直持有，没有调整", "记不清"},
        "Q25": {"没有明显不确定"},
        "Q26": {"目前不需要更多信息"},
    }
    return label in rules.get(qid, set())


def component_id(qid: str) -> str:
    """Return the schema-safe component ID while retaining the public item number."""
    return f"ITEM-{qid}"


def question_to_item(question: Question, variables: dict[str, str]) -> dict:
    qid = question.qid
    result = {
        "item_id": component_id(qid),
        "item_type": item_type(question),
        "prompt": question.prompt,
        "required": qid not in OPTIONAL_IDS,
        "research_question_ids": RQ_MAP.get(qid, []),
        "hypothesis_ids": HYP_MAP.get(qid, []),
        "randomize_options": False,
        "variable_name": variables[qid],
    }
    help_parts = [part for part in [question.instruction, question.skip_note] if part]
    if help_parts:
        result["help_text"] = "；".join(help_parts)
    fact_ids = FACT_MAP.get(qid, [])
    if fact_ids:
        result["quality_note"] = (
            "金融事实已映射至 FIN-FACT-FOF@0.1.0：" + "、".join(fact_ids)
        )
    elif qid == "Q32":
        result["quality_note"] = "仅用于本项目年龄分群；不形成正式客户标签。"
    elif qid == "Q9A":
        result["quality_note"] = "主观清晰度与客观知识题分开；记不清不占用中点。"

    if qid == "Q9A":
        result["scale"] = {
            "minimum": 1,
            "maximum": 5,
            "step": 1,
            "minimum_label": "非常不清楚",
            "maximum_label": "非常清楚",
            "direction": "LOW_TO_HIGH",
            "not_applicable_option": {
                "option_id": "OPT-Q9A-98",
                "label": "记不清",
                "value": 98,
            },
        }
    elif question.options:
        result["answer_options"] = [
            {
                "option_id": f"OPT-{qid}-{index:02d}",
                "label": label,
                "value": normalized_value(qid, label, index),
                **({"exclusive": True} if exclusive_option(qid, label) else {}),
            }
            for index, label in enumerate(question.options, start=1)
        ]
    return result


def analysis_role(qid: str) -> str:
    if qid == "C0":
        return "CONSENT"
    if qid == "S1":
        return "ELIGIBILITY"
    if qid == "Q32":
        return "SEGMENT"
    if qid in {"Q17", "Q18", "Q19", "Q20"}:
        return "KNOWLEDGE"
    if qid in {"Q11", "Q29"}:
        return "OPEN_TEXT"
    return "OUTCOME"


def output_variable(question: Question, variables: dict[str, str]) -> dict:
    qid = question.qid
    data_type = "ARRAY" if qid in MULTI_IDS else "INTEGER" if qid == "Q9A" else "STRING"
    result = {
        "variable_id": f"VAR-{qid}",
        "source_item_id": component_id(qid),
        "label": question.prompt,
        "data_type": data_type,
        "analysis_role": analysis_role(qid),
    }
    if qid == "C0":
        result["missing_value_rule"] = "缺失视为 CONSENT_NOT_GIVEN，不进入研究数据。"
    elif qid == "S1":
        result["missing_value_rule"] = "缺失视为甄别未完成；明确没有购买为 INELIGIBLE。"
    elif qid == "Q23":
        result["missing_value_rule"] = (
            "因 Q22 选择没有经历过或记不清而跳过时为 STRUCTURAL_MISSING；"
            "应展示但缺失时为 ITEM_MISSING。"
        )
    elif qid in OPTIONAL_IDS:
        result["missing_value_rule"] = "选答空白记为 ITEM_MISSING，但不据此判整份答卷无效。"
    else:
        result["missing_value_rule"] = (
            "应展示但缺失时记为 ITEM_MISSING；不确定、记不清或不愿回答按预设特殊值保留。"
        )

    value_mapping = {}
    special_values = {}
    for index, label in enumerate(question.options, start=1):
        value = normalized_value(qid, label, index)
        value_mapping[str(value)] = label
        semantic = value_semantic(label)
        if semantic:
            special_values[str(value)] = semantic
    if qid == "Q9A":
        value_mapping = {
            "1": "非常不清楚",
            "2": "比较不清楚",
            "3": "一般",
            "4": "比较清楚",
            "5": "非常清楚",
            "98": "记不清",
        }
        special_values = {"98": "UNKNOWN"}
    if value_mapping:
        result["value_mapping"] = value_mapping
    if special_values:
        result["special_value_meanings"] = special_values
    return result


def response_validity_plan() -> dict:
    return {
        "initial_status": "UNASSESSED",
        "final_statuses": ["VALID", "REVIEW_REQUIRED", "EXCLUDED"],
        "hard_exclusion_rules": [
            {
                "rule_id": "DQ-EX-CONSENT",
                "reason_code": "CONSENT_NOT_GIVEN",
                "condition": "C0 不同意或同意字段缺失。",
                "default_action": "EXCLUDE",
                "decision_authority": "确定性规则。",
            },
            {
                "rule_id": "DQ-EX-INELIGIBLE",
                "reason_code": "INELIGIBLE",
                "condition": "S1 明确选择没有实际购买过养老目标基金。",
                "default_action": "EXCLUDE",
                "decision_authority": "确定性甄别规则。",
            },
            {
                "rule_id": "DQ-EX-EMPTY",
                "reason_code": "NO_SUBSTANTIVE_DATA",
                "condition": "除同意与甄别外，全部核心模块均无正常作答或有效特殊值。",
                "default_action": "EXCLUDE",
                "decision_authority": "确定性完整度规则；保留原因。",
            },
            {
                "rule_id": "DQ-EX-SCHEMA",
                "reason_code": "SCHEMA_OR_VERSION_ERROR",
                "condition": "记录无法对应 InstrumentSpec 0.2.0 的字段或选项，且不能可靠修复。",
                "default_action": "EXCLUDE",
                "decision_authority": "数据责任人确认版本错误。",
            },
            {
                "rule_id": "DQ-EX-DUPLICATE",
                "reason_code": "CONFIRMED_DUPLICATE",
                "condition": "人工确认与另一记录重复，并按预设规则保留一份。",
                "default_action": "EXCLUDE",
                "decision_authority": "研究人员人工确认。",
            },
        ],
        "review_rules": [
            {
                "rule_id": "DQ-RV-ELIGIBILITY",
                "reason_code": "ELIGIBILITY_UNCERTAIN",
                "condition": "S1 选择不确定自己买的是不是养老目标基金。",
                "default_action": "REVIEW",
                "decision_authority": "研究人员结合 Q10/Q11 等自报信息复核；不得由模型猜测。",
            },
            {
                "rule_id": "DQ-RV-SPEED",
                "reason_code": "SPEEDING_SUSPECTED",
                "condition": "至少已有 30 份完成记录，且时长低于当批中位数三分之一。",
                "default_action": "REVIEW",
                "decision_authority": "研究人员结合其他信号复核；不得仅按时长排除。",
            },
            {
                "rule_id": "DQ-RV-CONFLICT",
                "reason_code": "LOGIC_CONFLICT",
                "condition": "关键事实或路由字段出现无法由题意和合法跳转解释的冲突。",
                "default_action": "REVIEW",
                "decision_authority": "研究人员复核原始记录。",
            },
            {
                "rule_id": "DQ-RV-DUPLICATE",
                "reason_code": "DUPLICATE_SUSPECTED",
                "condition": "匿名记录高度相似但尚未确认重复。",
                "default_action": "REVIEW",
                "decision_authority": "研究人员人工确认；不保留身份字段。",
            },
            {
                "rule_id": "DQ-RV-MISSING",
                "reason_code": "HIGH_ITEM_MISSINGNESS",
                "condition": "应展示的核心题目中 ITEM_MISSING 超过 20%。",
                "default_action": "REVIEW",
                "decision_authority": "研究人员判断剩余信息是否足以回答至少一个研究问题。",
            },
        ],
        "status_transition_rule": (
            "导入后先标 UNASSESSED；硬规则可转 EXCLUDED，其余异常只转 REVIEW_REQUIRED；"
            "人工解决全部复核后才冻结为 VALID 或 EXCLUDED。知识题答错、不确定、年龄不愿回答"
            "和开放题空白均不能单独触发排除。"
        ),
        "dataset_statuses": [
            "ANALYSIS_READY",
            "ANALYSIS_READY_WITH_LIMITS",
            "BLOCKED",
        ],
        "subgroup_reporting_rule": (
            "年龄组 n<10 不单独展示比例；10-19 标注小样本；n>=20 仍只作描述性比较，"
            "不推断总体、显著性或因果。"
        ),
    }


def build_instrument() -> dict:
    questions, section_order = parse_questionnaire(TARGET)
    variables = variable_map(TARGET)
    questions_by_section = {
        section: [q for q in questions if q.section == section] for section in section_order
    }
    sections = []
    for index, section_title in enumerate(section_order, start=1):
        section_questions = questions_by_section[section_title]
        section_rqs = sorted(
            {rq for question in section_questions for rq in RQ_MAP.get(question.qid, [])}
        )
        sections.append(
            {
                "section_id": f"SEC-{index:02d}",
                "title": section_title,
                "purpose": SECTION_PURPOSE[section_title],
                "research_question_ids": section_rqs,
                "randomize_items": False,
                "items": [
                    question_to_item(question, variables) for question in section_questions
                ],
            }
        )

    traceability = [
        {
            "component_id": component_id(question.qid),
            "component_type": "SURVEY_ITEM",
            "research_question_ids": RQ_MAP[question.qid],
            "hypothesis_ids": HYP_MAP.get(question.qid, []),
        }
        for question in questions
        if question.qid in RQ_MAP
    ]

    return {
        "metadata": {
            "schema_version": "0.2.0",
            "artifact_id": "INS-OTF-SURVEY-001",
            "artifact_type": "InstrumentSpec",
            "artifact_version": "0.2.0",
            "project_id": "PROJ-OTF-001",
            "title": "养老目标基金购买者的选择与持有认知问卷工具规范",
            "language": "zh-CN",
            "lifecycle_status": "DRAFT",
            "created_at": "2026-07-23T19:18:00+08:00",
            "updated_at": "2026-07-23T19:18:00+08:00",
            "created_by": {
                "actor_id": "AGENT-CAP-03",
                "actor_type": "AGENT",
                "role": "研究工具设计专家",
                "model_id": "codex",
                "capability_version": "0.5.0",
            },
            "upstream_refs": [
                {
                    "artifact_id": "RP-OTF-001",
                    "artifact_type": "ResearchPlan",
                    "artifact_version": "0.2.0",
                }
            ],
            "content_classification": "REAL",
            "sensitivity_level": "INTERNAL",
            "contains_personal_data": False,
            "change_summary": (
                "新增 S1 购买经历甄别、Q9A 五点信息清晰度、Q32 年龄段，"
                "并将变量值语义、答卷状态和数据集状态与问卷正文同版冻结。"
            ),
        },
        "research_plan_ref": {
            "artifact_id": "RP-OTF-001",
            "artifact_type": "ResearchPlan",
            "artifact_version": "0.2.0",
        },
        "instrument_type": "SURVEY",
        "instrument_mode": "STANDARD",
        "participant_facing_title": "您是怎么选择和持有养老基金的？",
        "purpose": (
            "描述养老目标基金实际购买者的购买路径、产品类型与下滑曲线认知、持有行为"
            "和信息需求；不进行产品推荐、风险测评或适当性判断。"
        ),
        "research_question_ids": [f"RQ-00{i}" for i in range(1, 7)],
        "hypothesis_ids": [f"HYP-00{i}" for i in range(1, 5)],
        "estimated_duration_minutes": 10,
        "languages": ["zh-CN"],
        "eligibility_rules": [
            {
                "rule_id": "ELIG-PURCHASE-YES",
                "criterion": "S1 明确表示目前持有或过去持有过养老目标基金。",
                "implementation_component_ids": [component_id("S1")],
                "outcome": "ELIGIBLE",
            },
            {
                "rule_id": "ELIG-PURCHASE-NO",
                "criterion": "S1 明确表示从未购买过养老目标基金。",
                "implementation_component_ids": [component_id("S1")],
                "outcome": "INELIGIBLE",
            },
        ],
        "consent_and_intro": {
            "participant_intro": (
                "本问卷了解养老目标基金购买者的选择、认知和持有信息需求，预计约 10 分钟。"
            ),
            "voluntary_statement": "参与完全自愿，可以随时停止。",
            "data_use_statement": (
                "答案只用于匿名用户研究，不用于投资建议、风险测评或产品推荐。"
            ),
            "withdrawal_statement": "提交前可关闭页面退出；不要求提供身份或账户信息。",
            "recording_notice": "本问卷不录制音频或视频。",
            "consent_capture": "CHECKBOX",
        },
        "test_materials": [],
        "learning_resources": [
            {
                "resource_id": "LR-IMA-OTF-001",
                "title": "易方达基金 IMA 知识库",
                "source_type": "IMA",
                "uri": IMA_URL,
                "purpose": "在核心答题结束后提供继续搜索养老目标基金金融常识的自愿学习入口。",
                "placement": "COMPLETION_PAGE",
                "recommended_search_prompt": (
                    "可搜索“养老目标基金”“目标日期基金”“目标风险基金”或“下滑曲线”。"
                ),
                "optional_for_participant": True,
                "evidence_use": "NOT_RESEARCH_EVIDENCE",
                "link_review_status": "UNREVIEWED",
                "review_notes": "发布前确认链接可访问；不记录点击、停留或身份关联。",
            }
        ],
        "material_measurement_plan": [],
        "traceability_map": traceability,
        "privacy_and_compliance": {
            "collects_personal_data": False,
            "personal_data_fields": [],
            "contact_data_separation": (
                "问卷不收集姓名、联系方式、账户号码、金额或持仓截图；外部触达名单不进入研究数据。"
            ),
            "financial_expression_review_required": True,
            "no_return_promise_confirmed": True,
            "suitability_data_separated": True,
            "review_notes": (
                "Q2、Q10、Q17—Q20 使用 FIN-FACT-FOF@0.1.0；正式发布仍需 Gate 2 人工审核。"
            ),
        },
        "pilot_config": {
            "required": False,
        },
        "survey_spec": {
            "sections": sections,
            "logic_rules": [
                {
                    "rule_id": "LOGIC-CONSENT-END",
                    "conditions": [
                        {
                            "source_item_id": component_id("C0"),
                            "operator": "EQUALS",
                            "value": "o02",
                        }
                    ],
                    "condition_join": "ALL",
                    "action": "TERMINATE",
                    "target_id": "END-CONSENT",
                    "priority": 1,
                },
                {
                    "rule_id": "LOGIC-INELIGIBLE-END",
                    "conditions": [
                        {
                            "source_item_id": component_id("S1"),
                            "operator": "EQUALS",
                            "value": "o02",
                        }
                    ],
                    "condition_join": "ALL",
                    "action": "TERMINATE",
                    "target_id": "END-INELIGIBLE",
                    "priority": 2,
                },
                {
                    "rule_id": "LOGIC-Q22-SKIP-Q23",
                    "conditions": [
                        {
                            "source_item_id": component_id("Q22"),
                            "operator": "EQUALS",
                            "value": "not_applicable",
                        }
                    ],
                    "condition_join": "ALL",
                    "action": "SKIP_TO",
                    "target_id": component_id("Q24"),
                    "priority": 3,
                },
            ],
            "quota_rules": [],
            "termination_messages": [
                {
                    "termination_id": "END-CONSENT",
                    "message": "感谢您的关注。您未同意参与，本问卷到此结束。",
                },
                {
                    "termination_id": "END-INELIGIBLE",
                    "message": "感谢您的关注。本研究面向实际购买过养老目标基金的投资者，本问卷到此结束。",
                },
            ],
            "output_variables": [
                output_variable(question, variables) for question in questions
            ],
            "response_validity_plan": response_validity_plan(),
            "progress_display": True,
            "back_navigation": False,
        },
    }


def validate_local(document_path: Path, instrument: dict) -> dict:
    questions, _ = parse_questionnaire(document_path)
    document = Document(document_path)
    question_ids = [question.qid for question in questions]
    if len(set(question_ids)) != len(question_ids):
        raise ValueError("Duplicate question IDs.")
    if "知识题答案与计分" not in "\n".join(p.text for p in document.paragraphs):
        raise ValueError("Internal answer appendix missing.")
    if "https://ima.qq.com/wiki/" not in document._element.xml:
        raise ValueError("IMA hyperlink field missing.")
    participant_text = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == "内部配置附录":
            break
        participant_text.append(paragraph.text)
    if "VALID" in "\n".join(participant_text):
        raise ValueError("Internal status language leaked into participant questionnaire.")
    return {
        "document": str(document_path),
        "questions": len(questions),
        "options": sum(len(question.options) for question in questions),
        "tables": len(document.tables),
        "instrument_sections": len(instrument["survey_spec"]["sections"]),
        "output_variables": len(instrument["survey_spec"]["output_variables"]),
    }


def main() -> None:
    update_word()
    instrument = build_instrument()
    INSTRUMENT_TARGET.write_text(
        json.dumps(instrument, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    summary = validate_local(TARGET, instrument)
    summary["instrument"] = str(INSTRUMENT_TARGET)
    print(json.dumps(summary, ensure_ascii=False))


if __name__ == "__main__":
    main()
